"""
PDF Parser — Phase 1 (Data Processing)  ·  Enhanced Edition
============================================================
Handles:
   Multi-column PDFs (X-axis aware column separation)
   Arabic RTL + reshaping + bidi + script normalisation  [Bonus Arabic]
   Scanned / image-only PDFs  →  OCR via pytesseract (graceful fallback)
   Table detection & extraction (structured text preserved)
   Header / footer detection and stripping
   Font-size-based section heading detection
   Word documents (.docx) support
   HTML documents parsing (BeautifulSoup)
   Noise cleaning with preserved structure
"""

import fitz # PyMuPDF
import re
import os
import io
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple

# ── Optional: Arabic ──────────────────────────────────────────────────────────
try:
    import arabic_reshaper
    from bidi.algorithm import get_display
    ARABIC_SUPPORT = True
except ImportError:
    ARABIC_SUPPORT = False
    print("[Parser] ⚠️  arabic-reshaper / python-bidi not installed — Arabic reshaping disabled.")

# ── Optional: OCR (scanned PDFs) ─────────────────────────────────────────────
try:
    import pytesseract
    from PIL import Image
    OCR_SUPPORT = True
except ImportError:
    OCR_SUPPORT = False
    print("[Parser] ⚠️  pytesseract / Pillow not installed — OCR disabled.")

# ── Optional: DOCX support ────────────────────────────────────────────────────
try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("[Parser] ⚠️  python-docx not installed — .docx parsing disabled.")

# ── Optional: HTML support ────────────────────────────────────────────────────
try:
    from bs4 import BeautifulSoup
    HTML_SUPPORT = True
except ImportError:
    HTML_SUPPORT = False
    print("[Parser] ⚠️  beautifulsoup4 not installed — HTML parsing disabled.")


# ─────────────────────────────────────────────────────────────────────────────
#  Regex catalogue
# ─────────────────────────────────────────────────────────────────────────────
_URL_RE     = re.compile(r"https?://\S+|www\.\S+")
_WS_RE      = re.compile(r"[ \t]{2,}")
_NEWLINE_RE = re.compile(r"\n{3,}")
_NOISE_RE   = re.compile(
    r"[^\w\s\u0600-\u06FF\u0750-\u077F\u200c\u200d.,;:()\-+@/#%&'\"\!?،؟]"
)

_ARABIC_NORM = str.maketrans({
    "أ": "ا", "إ": "ا", "آ": "ا", "ٱ": "ا",
    "ة": "ه",
    "ى": "ي", "ئ": "ي",
    "ؤ": "و",
    "\u0640": "",
    "\u064B": "", "\u064C": "", "\u064D": "",   # tanwin
    "\u064E": "", "\u064F": "", "\u0650": "",   # fatha / damma / kasra
    "\u0651": "", "\u0652": "",                 # shadda / sukun
})


# ─────────────────────────────────────────────────────────────────────────────
#  Helper: column-aware block sorting
# ─────────────────────────────────────────────────────────────────────────────

def _sort_blocks_column_aware(
    blocks: List[Tuple], page_width: float, col_threshold: float = 0.45
) -> List[Tuple]:
   
    left_col, right_col, uncertain = [], [], []

    for b in blocks:
        x_centre = (b[0] + b[2]) / 2
        if x_centre < page_width * col_threshold:
            left_col.append(b)
        elif x_centre > page_width * (1 - col_threshold):
            right_col.append(b)
        else:
            uncertain.append(b)

    if not left_col or not right_col:
        return sorted(blocks, key=lambda b: (b[1], b[0]))

    left_col.sort(key=lambda b: b[1])
    right_col.sort(key=lambda b: b[1])
    uncertain.sort(key=lambda b: (b[1], b[0]))

    return left_col + right_col + uncertain


# ─────────────────────────────────────────────────────────────────────────────
#  Helper: header / footer stripping
# ─────────────────────────────────────────────────────────────────────────────

def _is_header_footer(y: float, page_height: float, margin_ratio: float = 0.06) -> bool:
    """Returns True if the block sits in the top or bottom margin zone."""
    margin = page_height * margin_ratio
    return y < margin or y > (page_height - margin)


# ─────────────────────────────────────────────────────────────────────────────
#  Helper: table extraction
# ─────────────────────────────────────────────────────────────────────────────

def _extract_tables_from_page(page: fitz.Page) -> str:
    """Uses PyMuPDF's find_tables() (≥ 1.23). Falls back to empty string."""
    try:
        tabs = page.find_tables()
        if not tabs.tables:
            return ""
        lines = []
        for table in tabs.tables:
            for row in table.extract():
                cells = [str(c).strip() if c else "" for c in row]
                lines.append(" | ".join(cells))
            lines.append("")
        return "\n".join(lines)
    except AttributeError:
        return ""


# ─────────────────────────────────────────────────────────────────────────────
#  Helper: font-size → section headings
# ─────────────────────────────────────────────────────────────────────────────

def _extract_headings(page: fitz.Page, body_font_size_threshold: float = 13.0) -> List[str]:
    """Returns text spans whose font size exceeds the threshold — likely section headings."""
    headings = []
    for block in page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]:
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                if span.get("size", 0) >= body_font_size_threshold:
                    txt = span["text"].strip()
                    if txt and len(txt) > 2:
                        headings.append(txt)
    return headings


# ─────────────────────────────────────────────────────────────────────────────
#  Helper: OCR fallback for image-based pages
# ─────────────────────────────────────────────────────────────────────────────

def _ocr_page(page: fitz.Page, lang: str = "ara+eng") -> str:
    """Rasterises the page at 200 DPI and runs Tesseract OCR."""
    if not OCR_SUPPORT:
        return ""
    try:
        mat = fitz.Matrix(200 / 72, 200 / 72)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        return pytesseract.image_to_string(img, lang=lang)
    except Exception as exc:
        print(f"[Parser] ⚠️  OCR failed on page {page.number + 1}: {exc}")
        return ""


def _page_has_real_text(page: fitz.Page, min_chars: int = 30) -> bool:
    """Quick check: does the page carry enough selectable text, or is it a scan?"""
    return len(page.get_text("text").strip()) >= min_chars


# ─────────────────────────────────────────────────────────────────────────────
#  Main Parser class
# ─────────────────────────────────────────────────────────────────────────────

class PDFParser:
   
    # ── Arabic utilities ─────────────────────────────────────────────────────

    @staticmethod
    def _arabic_ratio(text: str) -> float:
        if not text:
            return 0.0
        arabic = sum(1 for c in text if "\u0600" <= c <= "\u06FF")
        return arabic / max(len(text), 1)

    @staticmethod
    def _normalise_arabic(text: str) -> str:
       
        return text.translate(_ARABIC_NORM)

    def _fix_arabic(self, text: str) -> str:
        """Normalise → reshape → bidi reorder for correct RTL rendering."""
        text = self._normalise_arabic(text)
        if not ARABIC_SUPPORT:
            return text
        try:
            reshaped = arabic_reshaper.reshape(text)
            return get_display(reshaped)
        except Exception:
            return text

    # ── Cleaning ─────────────────────────────────────────────────────────────

    def _clean(self, text: str) -> str:
        text = _URL_RE.sub(" ", text)
        text = _NOISE_RE.sub(" ", text)
        text = _WS_RE.sub(" ", text)
        text = _NEWLINE_RE.sub("\n\n", text)
        return text.strip()

    # ── Single PDF ────────────────────────────────────────────────────────────

    def parse_pdf(self, file_path: str) -> Dict[str, Any]:
        
        path = Path(file_path)
        doc  = fitz.open(str(path))
        pages_count = len(doc)

        raw_blocks:   List[str] = []
        table_texts:  List[str] = []
        all_headings: List[str] = []
        scanned_pages = 0

        for page in doc:
            page_width  = page.rect.width
            page_height = page.rect.height

            table_text = _extract_tables_from_page(page)
            if table_text:
                table_texts.append(table_text)

            if not _page_has_real_text(page):
                scanned_pages += 1
                ocr_text = _ocr_page(page)
                if ocr_text.strip():
                    raw_blocks.append(ocr_text)
                continue

            all_headings.extend(_extract_headings(page))

            blocks = page.get_text("blocks", sort=False)
            text_blocks = [
                b for b in blocks
                if b[6] == 0
                and b[4].strip()
                and not _is_header_footer(b[1], page_height)
            ]
            sorted_blocks = _sort_blocks_column_aware(text_blocks, page_width)
            raw_blocks.extend(b[4] for b in sorted_blocks)

        doc.close()

        raw_text = "\n".join(raw_blocks)
        if table_texts:
            raw_text += "\n\n[TABLES]\n" + "\n".join(table_texts)

        arabic_ratio = self._arabic_ratio(raw_text)
        if arabic_ratio > 0.15:
            raw_text = self._fix_arabic(raw_text)
            language = "arabic" if arabic_ratio > 0.5 else "mixed"
        else:
            language = "english"

        cleaned = self._clean(raw_text)

        seen: set = set()
        unique_headings = []
        for h in all_headings:
            key = h.lower().strip()
            if key not in seen and len(key) > 2:
                seen.add(key)
                unique_headings.append(h)

        return {
            "filename":     path.name,
            "stem":         path.stem,
            "text":         cleaned,
            "pages":        pages_count,
            "language":     language,
            "arabic_ratio": round(arabic_ratio, 3),
            "source":       str(path),
            "size_bytes":   os.path.getsize(str(path)),
            "headings":     unique_headings[:20],
        }

    # ── DOCX support ─────────────────────────────────────────────────────────

    def parse_docx(self, file_path: str) -> Dict[str, Any]:
        """Parse a Word document (.docx) using python-docx."""
        if not DOCX_SUPPORT:
            raise RuntimeError(
                "python-docx is not installed. Run: pip install python-docx"
            )

        path = Path(file_path)
        doc  = DocxDocument(str(path))

        paragraphs: List[str] = []
        headings:   List[str] = []

        for para in doc.paragraphs:
            txt = para.text.strip()
            if not txt:
                continue
            if para.style.name.startswith("Heading"):
                headings.append(txt)
            paragraphs.append(txt)

        table_lines: List[str] = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_lines.append(" | ".join(cells))
            table_lines.append("")

        raw_text = "\n".join(paragraphs)
        if table_lines:
            raw_text += "\n\n[TABLES]\n" + "\n".join(table_lines)

        arabic_ratio = self._arabic_ratio(raw_text)
        if arabic_ratio > 0.15:
            raw_text = self._fix_arabic(raw_text)
            language = "arabic" if arabic_ratio > 0.5 else "mixed"
        else:
            language = "english"

        cleaned = self._clean(raw_text)

        return {
            "filename":     path.name,
            "stem":         path.stem,
            "text":         cleaned,
            "pages":        len(doc.sections),
            "language":     language,
            "arabic_ratio": round(arabic_ratio, 3),
            "source":       str(path),
            "size_bytes":   os.path.getsize(str(path)),
            "headings":     headings[:20],
        }

    # ── HTML support ─────────────────────────────────────────────────────────

    def parse_html(self, file_path: str) -> Dict[str, Any]:
        
        if not HTML_SUPPORT:
            raise RuntimeError(
                "beautifulsoup4 is not installed. Run: pip install beautifulsoup4"
            )

        path = Path(file_path)
        raw_html = path.read_text(encoding="utf-8", errors="replace")
        soup = BeautifulSoup(raw_html, "html.parser")

        # Strip noise tags
        for tag in soup(["script", "style", "nav", "footer", "header", "meta", "link"]):
            tag.decompose()

        # Extract headings
        headings: List[str] = []
        for h_tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6"]):
            txt = h_tag.get_text(strip=True)
            if txt and len(txt) > 2:
                headings.append(txt)

        # Extract tables (and remove from soup so they don't repeat in body text)
        table_lines: List[str] = []
        for table in soup.find_all("table"):
            for row in table.find_all("tr"):
                cells = [td.get_text(strip=True) for td in row.find_all(["td", "th"])]
                if any(cells):
                    table_lines.append(" | ".join(cells))
            table_lines.append("")
            table.decompose()

        raw_text = soup.get_text(separator="\n")
        if table_lines:
            raw_text += "\n\n[TABLES]\n" + "\n".join(table_lines)

        arabic_ratio = self._arabic_ratio(raw_text)
        if arabic_ratio > 0.15:
            raw_text = self._fix_arabic(raw_text)
            language = "arabic" if arabic_ratio > 0.5 else "mixed"
        else:
            language = "english"

        cleaned = self._clean(raw_text)

        return {
            "filename":     path.name,
            "stem":         path.stem,
            "text":         cleaned,
            "pages":        1,
            "language":     language,
            "arabic_ratio": round(arabic_ratio, 3),
            "source":       str(path),
            "size_bytes":   os.path.getsize(str(path)),
            "headings":     headings[:20],
        }

    # ── Universal entry point ─────────────────────────────────────────────────

    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """
        Route to the correct parser based on file extension.
        Supports: .pdf | .docx | .html | .htm
        """
        ext = Path(file_path).suffix.lower()
        if ext == ".pdf":
            return self.parse_pdf(file_path)
        elif ext == ".docx":
            return self.parse_docx(file_path)
        elif ext in (".html", ".htm"):
            return self.parse_html(file_path)
        else:
            raise ValueError(
                f"Unsupported file type: '{ext}'. Supported: .pdf, .docx, .html, .htm"
            )

    # ── Directory batch parsing ───────────────────────────────────────────────

    def parse_directory(
        self,
        dir_path: str,
        extensions: Optional[List[str]] = None,
    ) -> List[Dict[str, Any]]:
       
        if extensions is None:
            extensions = [".pdf", ".docx", ".html", ".htm"]

        dir_path = Path(dir_path)
        files: List[Path] = []
        for ext in extensions:
            files.extend(sorted(dir_path.glob(f"*{ext}")))

        if not files:
            print(f"[Parser] ⚠️  No supported files found in: {dir_path}")
            return []

        docs, skipped = [], []
        for f in files:
            try:
                doc = self.parse_file(str(f))
                docs.append(doc)
                print(
                    f"[Parser] ✅ {f.name} | "
                    f"{doc['pages']}p | "
                    f"{doc['language']} | "
                    f"{doc['arabic_ratio']:.0%} Arabic"
                )
            except Exception as exc:
                print(f"[Parser] ❌ {f.name}: {exc}")
                skipped.append(f.name)

        print(f"[Parser] Done — {len(docs)} parsed, {len(skipped)} skipped.")
        return docs